import json
from pathlib import Path

import docx  # python-docx
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from config import UPLOAD_DIR


def _insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    """
    Insert a new paragraph with `text` immediately after `paragraph`.
    """
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.text = text
    return new_para


def _apply_changes_to_doc(doc: docx.Document, changes) -> None:
    # Separate actions so index shifts are handled safely
    modifies = [c for c in changes if c.get("action") == "modify"]
    removes = sorted(
        (c for c in changes if c.get("action") == "remove"),
        key=lambda c: c.get("line_number", 0),
        reverse=True,
    )
    adds = sorted(
        (c for c in changes if c.get("action") == "add_new_line"),
        key=lambda c: c.get(
            "after_line_number",
            c.get("after_line", c.get("line_number", 0)),
        ),
        reverse=True,
    )

    # Apply modifications (safe in any order)
    for change in modifies:
        line_number = change.get("line_number")
        new_text = change.get("new_text", "")
        if not isinstance(line_number, int):
            continue
        paragraphs = doc.paragraphs
        idx = line_number - 1  # 1-based to 0-based
        if 0 <= idx < len(paragraphs):
            para = paragraphs[idx]
            # Preserve basic run formatting (bold/italic/color/font) by
            # cloning the first run's formatting onto the new text.
            if para.runs:
                first_run = para.runs[0]
                # Clear existing runs
                for r in list(para.runs):
                    r.clear()
                new_run = para.add_run(new_text)
                new_run.bold = first_run.bold
                new_run.italic = first_run.italic
                new_run.underline = first_run.underline
                new_run.style = first_run.style
                if first_run.font is not None and new_run.font is not None:
                    new_run.font.name = first_run.font.name
                    new_run.font.size = first_run.font.size
                    new_run.font.bold = first_run.font.bold
                    new_run.font.italic = first_run.font.italic
                    new_run.font.underline = first_run.font.underline
                    new_run.font.color.rgb = getattr(first_run.font.color, "rgb", None)
            else:
                para.text = new_text

    # Apply removals from bottom to top so indices above are unaffected
    for change in removes:
        line_number = change.get("line_number")
        if not isinstance(line_number, int):
            continue
        paragraphs = doc.paragraphs
        idx = line_number - 1
        if 0 <= idx < len(paragraphs):
            p = paragraphs[idx]
            p._p.getparent().remove(p._p)

    # Apply additions from bottom to top so earlier indices stay stable
    for change in adds:
        # Support both "after_line_number" and "after_line" keys
        after_line_number = change.get(
            "after_line_number",
            change.get("after_line", change.get("line_number")),
        )
        new_text = change.get("new_text", "")
        if not isinstance(after_line_number, int):
            continue
        paragraphs = doc.paragraphs
        idx = after_line_number - 1
        if 0 <= idx < len(paragraphs):
            _insert_paragraph_after(paragraphs[idx], new_text)
        elif idx == len(paragraphs):
            # If asked to add after the last line, just append
            doc.add_paragraph(new_text)


def apply_changes_in_memory(file_path: Path, changes) -> bytes:
    """
    Open the given DOCX, apply changes, and return the updated
    document as bytes in memory (no files written to disk).
    """
    if not file_path.exists():
        raise FileNotFoundError(f"Source DOCX not found: {file_path}")

    doc = docx.Document(str(file_path))
    _apply_changes_to_doc(doc, changes)

    from io import BytesIO

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


if __name__ == "__main__":
    out = apply_changes_from_json()
    print(f"Saved updated DOCX to: {out}")
